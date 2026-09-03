"""Формирование отчётов: Telegram (HTML-разметка) и .docx (v1.1.0, для юристов/HR)."""

from __future__ import annotations

import html
from datetime import datetime, timezone
from io import BytesIO

from secretary.llm.deepseek import MeetingSummary
from secretary.stt.base import TranscriptResult


def _esc(text: str) -> str:
    return html.escape(text or "", quote=False)


def render_tg_html(tr: TranscriptResult, summary: MeetingSummary, order_id: int) -> str:
    """Отчёт для Telegram в HTML-разметке (паттерн tg-mcp)."""
    lines: list[str] = []
    lines.append(f"<b>📋 Итоги созвона #{order_id}</b>")
    lines.append(f"⏱ Длительность: {_fmt_duration(tr.audio_duration_sec)} · Спикеры: {'✓' if tr.with_speakers else '—'}")
    lines.append("")
    lines.append("<b>Саммари</b>")
    lines.append(_esc(summary.summary) or "—")
    if summary.key_topics:
        lines.append("")
        lines.append("<b>Темы</b>")
        lines.append(" · ".join(_esc(t) for t in summary.key_topics[:8]))
    if summary.decisions:
        lines.append("")
        lines.append("<b>✅ Решения</b>")
        lines.extend(f"• {_esc(d)}" for d in summary.decisions)
    if summary.tasks:
        lines.append("")
        lines.append("<b>Задачи</b>")
        for t in summary.tasks[:12]:
            prio = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(str(t.get("priority", "")).lower(), "⚪")
            owner = f" → {_esc(str(t.get('owner', '')))}" if t.get("owner") else ""
            lines.append(f"{prio} {_esc(str(t.get('task', '')))}<i>{owner}</i>")
    if tr.utterances:
        lines.append("")
        lines.append("<b>Фрагменты</b>")
        # первые 10 реплик — вкус к разговору (полная стенограмма — в .docx)
        for u in tr.utterances[:10]:
            lines.append(f"<b>Спикер {_esc(u.speaker)}:</b> {_esc(u.text)}")
        if len(tr.utterances) > 10:
            lines.append(f"<i>… ещё {len(tr.utterances) - 10} реплик — полная версия в .docx</i>")
    lines.append("")
    lines.append("<i>Расшифровка предоставленного файла · согласие на запись — ответственность клиента</i>")
    return "\n".join(lines)


def render_txt(tr: TranscriptResult) -> str:
    """Полная стенограмма текстом (для .docx и скачивания)."""
    if tr.utterances:
        return "\n".join(f"[Спикер {u.speaker}] {u.text}" for u in tr.utterances)
    return tr.text


def render_docx(tr: TranscriptResult, summary: MeetingSummary, order_id: int) -> BytesIO:
    """Отчёт .docx (v1.1.0) — вид «документа» для юристов/HR."""
    try:
        from docx import Document  # noqa: PLC0415
        from docx.shared import Pt  # noqa: PLC0415
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("python-docx не установлен (pip install python-docx)") from e

    doc = Document()
    doc.add_heading(f"Итоги созвона #{order_id}", level=0)
    stamp = datetime.now(timezone.utc).strftime("%d.%m.%Y %H:%M UTC")
    doc.add_paragraph(f"Сформировано: {stamp} · Длительность: {_fmt_duration(tr.audio_duration_sec)}")

    doc.add_heading("Саммари", level=1)
    doc.add_paragraph(summary.summary or "—")

    if summary.key_topics:
        doc.add_heading("Ключевые темы", level=1)
        for t in summary.key_topics:
            doc.add_paragraph(t, style="List Bullet")

    if summary.decisions:
        doc.add_heading("Решения", level=1)
        for d in summary.decisions:
            doc.add_paragraph(d, style="List Bullet")

    if summary.tasks:
        doc.add_heading("Задачи", level=1)
        for t in summary.tasks:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(t.get("task", "")))
            if t.get("owner"):
                p.add_run(f" — {t['owner']}")
            p.add_run(f" [{t.get('priority', '')}]")

    doc.add_heading("Стенограмма", level=1)
    for u in tr.utterances:
        p = doc.add_paragraph()
        run = p.add_run(f"Спикер {u.speaker}: ")
        run.bold = True
        p.add_run(u.text)
    if not tr.utterances and tr.text:
        doc.add_paragraph(tr.text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _fmt_duration(sec: float) -> str:
    sec = int(sec or 0)
    h, rem = divmod(sec, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h} ч {m:02d} мин"
    if m:
        return f"{m} мин {s:02d} с"
    return f"{s} с"
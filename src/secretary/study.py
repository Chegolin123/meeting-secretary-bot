"""Режим «Личный инструмент: конспекты пар/лекций» (study_mode).

Отличается от продукта-секретаря: вместо решений/задач — учебный конспект:
тема/предмет, ключевые тезисы, определения, формулы, вопросы для повторения.
Включается STUDY_MODE=true (или --study у моста). Учебные поля кладутся
в summary_json тем же контрактом — кабинет и мост не ломаются.
"""

from __future__ import annotations

import dataclasses
import json
import re

from secretary.stt.base import TranscriptResult

STUDY_SYSTEM_PROMPT = (
    'Ты — ассистент для конспектирования лекций. По транскрипту пары/лекции '
    'составь учебный конспект. Отвечай ТОЛЬКО валидным JSON без markdown-разметки:\n'
    '{"subject": "название предмета (если ясно из текста, иначе пустая строка)", '
    '"lecture_topic": "тема лекции", '
    '"key_points": ["главные тезисы"], '
    '"definitions": [{"term": "термин", "definition": "определение"}], '
    '"formulas": ["формулы/код/важные детали — если были"], '
    '"review_questions": ["3-7 вопросов для повторения"], '
    '"unclear": ["что осталось непонятным / стоит уточнить"]}\n'
    'Не выдумывай: только то, что было в лекции. Если чего-то не было — пустой список.'
)


@dataclasses.dataclass
class LectureSummary:
    subject: str
    lecture_topic: str
    key_points: list[str]
    definitions: list[dict]
    formulas: list[str]
    review_questions: list[str]
    unclear: list[str]

    def to_dict(self) -> dict:
        return {
            'subject': self.subject,
            'lecture_topic': self.lecture_topic,
            'key_points': self.key_points,
            'definitions': self.definitions,
            'formulas': self.formulas,
            'review_questions': self.review_questions,
            'unclear': self.unclear,
        }


def build_study_system_prompt() -> str:
    return STUDY_SYSTEM_PROMPT


def parse_lecture(content: str) -> LectureSummary:
    """Парсинг JSON-ответа модели в конспект; не-JSON → пустой каркас (не выдумываем)."""
    text = content.strip()
    text = re.sub(r'^```(?:json)?\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find('{'), text.rfind('}')
        if start == -1 or end == -1:
            return LectureSummary('', '', [], [], [], [], ['не удалось разобрать ответ модели'])
        data = json.loads(text[start : end + 1])

    def lst(key: str) -> list:
        v = data.get(key) or []
        return [str(x) for x in v if str(x).strip()] if isinstance(v, list) else []

    def lst_of_dict(key: str) -> list:
        v = data.get(key) or []
        out = []
        for x in v if isinstance(v, list) else []:
            if isinstance(x, dict) and str(x.get('term', '')).strip():
                out.append({'term': str(x['term']), 'definition': str(x.get('definition', ''))})
            elif isinstance(x, str) and x.strip():
                out.append({'term': x.strip(), 'definition': ''})
        return out

    return LectureSummary(
        subject=str(data.get('subject', '')).strip(),
        lecture_topic=str(data.get('lecture_topic', '')).strip(),
        key_points=lst('key_points'),
        definitions=lst_of_dict('definitions'),
        formulas=lst('formulas'),
        review_questions=lst('review_questions'),
        unclear=lst('unclear'),
    )


def _esc(text: str) -> str:
    return html_escape(text)


def html_escape(text: str) -> str:
    import html  # noqa: PLC0415

    return html.escape(text or '', quote=False)


def _fmt_duration(sec: float) -> str:
    sec = int(sec or 0)
    m, s = divmod(sec, 60)
    return f'{m} мин {s:02d} с' if m else f'{s} с'


def render_tg_lecture(tr: TranscriptResult, lecture: LectureSummary, order_id: int) -> str:
    """Учебный отчёт в Telegram (HTML)."""
    lines = [f"<b>🎓 Конспект #{order_id}</b>"]
    if lecture.subject:
        lines.append(f"📚 <b>Предмет:</b> {_esc(lecture.subject)}")
    if lecture.lecture_topic:
        lines.append(f"📌 <b>Тема:</b> {_esc(lecture.lecture_topic)}")
    lines.append(f"⏱ Длительность: {_fmt_duration(tr.audio_duration_sec)} · Спикеров: {'✓' if tr.with_speakers else '—'}")
    lines.append('')
    lines.append('<b>Ключевые тезисы</b>')
    lines += [f"• {_esc(p)}" for p in lecture.key_points] or ['—']
    if lecture.definitions:
        lines.append('')
        lines.append('<b>Определения</b>')
        lines += [f"• <b>{_esc(d['term'])}</b> — {_esc(d['definition'])}" for d in lecture.definitions]
    if lecture.formulas:
        lines.append('')
        lines.append('<b>Формулы / важные детали</b>')
        lines += [f"• {_esc(f)}" for f in lecture.formulas]
    if lecture.review_questions:
        lines.append('')
        lines.append('<b>Вопросы для повторения</b>')
        lines += [f"{i + 1}. {_esc(q)}" for i, q in enumerate(lecture.review_questions)]
    if lecture.unclear:
        lines.append('')
        lines.append('<i>📌 Уточнить:</i> ' + '; '.join(_esc(u) for u in lecture.unclear))
    lines.append('')
    lines.append('<i>Конспект по предоставленной записи лекции</i>')
    return '\n'.join(lines)


def render_docx_lecture(tr: TranscriptResult, lecture: LectureSummary, order_id: int):
    """Учебный конспект в .docx (для печати/чтения)."""
    from io import BytesIO  # noqa: PLC0415

    from docx import Document  # noqa: PLC0415

    doc = Document()
    doc.add_heading(f'Конспект лекции #{order_id}', level=0)
    if lecture.subject:
        doc.add_paragraph(f'Предмет: {lecture.subject}')
    if lecture.lecture_topic:
        doc.add_paragraph(f'Тема: {lecture.lecture_topic}')

    def section(title: str, items: list[str]) -> None:
        if not items:
            return
        doc.add_heading(title, level=1)
        for it in items:
            doc.add_paragraph(it, style='List Bullet')

    section('Ключевые тезисы', lecture.key_points)
    if lecture.definitions:
        doc.add_heading('Определения', level=1)
        for d in lecture.definitions:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(str(d['term']))
            run.bold = True
            p.add_run(f" — {d['definition']}")
    section('Формулы / важные детали', lecture.formulas)
    section('Вопросы для повторения', lecture.review_questions)

    doc.add_heading('Стенограмма', level=1)
    for u in tr.utterances:
        p = doc.add_paragraph()
        run = p.add_run(f'Спикер {u.speaker}: ')
        run.bold = True
        p.add_run(u.text)
    if not tr.utterances and tr.text:
        doc.add_paragraph(tr.text)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def lecture_to_summary_json(lecture: LectureSummary) -> str:
    return json.dumps(lecture.to_dict(), ensure_ascii=False)